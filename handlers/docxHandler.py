import os
import re
import warnings

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Cm, Pt
from docx.text.run import Run

from .contentHandler import ANSWER_PATTERNS, IMAGE_FORMATS, txt_to_json


def create_element(name):
    return OxmlElement(name)


def create_attribute(element: BaseOxmlElement, name: str, value: str):
    element.set(ns.qn(name), value)


def add_page_number(run: Run):
    page_element = create_element("w:t")
    create_attribute(page_element, "xml:space", "preserve")
    page_element.text = "Trang "

    run._r.append(page_element)

    fldChar1 = create_element("w:fldChar")
    create_attribute(fldChar1, "w:fldCharType", "begin")

    instrText1 = create_element("w:instrText")
    create_attribute(instrText1, "xml:space", "preserve")
    instrText1.text = "PAGE"

    fldChar2 = create_element("w:fldChar")
    create_attribute(fldChar2, "w:fldCharType", "end")

    run._r.append(fldChar1)
    run._r.append(instrText1)
    run._r.append(fldChar2)

    slash = create_element("w:t")
    create_attribute(slash, "xml:space", "preserve")
    slash.text = " / "

    run._r.append(slash)

    fldChar3 = create_element("w:fldChar")
    create_attribute(fldChar3, "w:fldCharType", "begin")

    instrText2 = create_element("w:instrText")
    create_attribute(instrText2, "xml:space", "preserve")
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element("w:fldChar")
    create_attribute(fldChar4, "w:fldCharType", "end")

    run._r.append(fldChar3)
    run._r.append(instrText2)
    run._r.append(fldChar4)


def docx_handler(dir_input: str, dir_output: str):
    with open(os.path.join(dir_input), "r", encoding="utf-8") as f:
        content = f.read()

    problems = txt_to_json(content)

    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        footer = section.footer
        footer_para = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )
        add_page_number(footer_para.add_run())
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    n_problems_of_ptype = {}

    my_table = doc.add_table(0, 2)

    for problem in problems.values():
        answer_prefix, answer_content = problem["answers"][0]
        ptype = ""
        if re.match(ANSWER_PATTERNS["mctf"], answer_prefix):
            ptype = "multiple_choice"
        if re.match(ANSWER_PATTERNS["shortans"], answer_prefix) and re.match(
            r"^[SD]+$", answer_content
        ):
            ptype = "true_false"

        if not ptype:
            warnings.warn("Skipping problem of unknown or unsupported type")
            continue

        if ptype in n_problems_of_ptype.keys():
            n_problems_of_ptype[ptype] += 1
        else:
            n_problems_of_ptype[ptype] = 1

        question = problem["question"]
        answers = problem["answers"]
        solution = problem["solution"]
        medias = problem["medias"]

        if ptype == "multiple_choice":
            row = my_table.add_row()
            row.cells[0].add_paragraph().add_run(
                f"Câu 1.{n_problems_of_ptype[ptype]}"
            ).bold = True
            r = row.cells[1].add_paragraph()
            r.add_run(question)
            for media_path in medias:
                media_format = os.path.splitext(media_path)[1][1:]
                if media_format in IMAGE_FORMATS:
                    r.add_run("\n\n")
                    r.add_run().add_picture(media_path, width=Cm(5.0))

            answer = 0
            for i, (prefix, choice) in enumerate(answers):
                if prefix[0] == "*":
                    answer = i
                row = my_table.add_row()
                row.cells[0].add_paragraph().add_run(chr(ord("A") + i)).bold = True
                row.cells[1].add_paragraph().add_run(choice)

            row = my_table.add_row()
            row.cells[0].add_paragraph().add_run("Lời giải").bold = True
            row.cells[1].add_paragraph().add_run(solution)

            row = my_table.add_row()
            row.cells[0].add_paragraph().add_run("Đáp án").bold = True
            row.cells[1].add_paragraph().add_run(chr(ord("A") + answer))

        if ptype == "true_false":
            question_without_statements, *statements = [
                chunk.strip()
                for chunk in re.split(
                    r"^\*?[a-zA-z]\)",
                    "\n".join([line.strip() for line in question.splitlines()]),
                    flags=re.MULTILINE,
                )
                if chunk
            ]

            row = my_table.add_row()
            row.cells[0].add_paragraph().add_run(
                f"Câu 2.{n_problems_of_ptype[ptype]}"
            ).bold = True
            r = row.cells[1].add_paragraph()
            r.add_run(question_without_statements)
            for media_path in medias:
                media_format = os.path.splitext(media_path)[1][1:]
                if media_format in IMAGE_FORMATS:
                    r.add_run("\n\n")
                    r.add_run().add_picture(media_path, width=Cm(5.0))

            answer = 0
            for i, statement in enumerate(statements):
                row = my_table.add_row()
                row.cells[0].add_paragraph().add_run(
                    f"{chr(ord('A') + i)})"
                ).bold = True
                row.cells[1].add_paragraph().add_run(statement)

            row = my_table.add_row()
            row.cells[0].add_paragraph().add_run("Lời giải").bold = True
            row.cells[1].add_paragraph().add_run(solution)

            row = my_table.add_row()
            row.cells[0].add_paragraph().add_run("Đáp án").bold = True
            row.cells[1].add_paragraph().add_run(answers[0][1])

    doc.save(os.path.join(dir_output, "exam.docx"))
